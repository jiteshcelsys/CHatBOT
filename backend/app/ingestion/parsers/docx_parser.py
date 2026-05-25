"""
DOCX parser — extracts text from paragraphs and tables via python-docx.

Why not use LangChain's Docx2txtLoader?
  It drops table content and header/footer text.  Direct python-docx access
  gives us full control: paragraphs, headings, and table cells are all captured.
"""
import logging
from pathlib import Path

import docx
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _extract_table_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


async def parse_docx(file_path: str) -> list[Document]:
    path = Path(file_path)
    doc = docx.Document(str(path))

    sections: list[str] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            # Paragraph or heading
            para = docx.text.paragraph.Paragraph(block, doc)
            text = para.text.strip()
            if text:
                sections.append(text)

        elif tag == "tbl":
            # Table — convert to pipe-delimited text
            table = docx.table.Table(block, doc)
            table_text = _extract_table_text(table)
            if table_text:
                sections.append(table_text)

    full_text = "\n\n".join(sections)

    result = Document(
        page_content=full_text,
        metadata={
            "source": path.name,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )
    logger.info(
        "DOCX parsed | file=%s paragraphs=%d tables=%d chars=%d",
        path.name, len(doc.paragraphs), len(doc.tables), len(full_text),
    )
    return [result]
